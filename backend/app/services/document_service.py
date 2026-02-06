import os
import re
import uuid
import logging
from html.parser import HTMLParser
from io import StringIO
from sqlalchemy.orm import Session
from fastapi import UploadFile
import PyPDF2
from docx import Document as DocxDocument

from ..models import Document
from ..config import get_settings
from ..exceptions import FileUploadError, NotFoundError

logger = logging.getLogger(__name__)
settings = get_settings()


class HTMLTextExtractor(HTMLParser):
    """Extract text from HTML while preserving basic structure."""
    
    def __init__(self):
        super().__init__()
        self.result = StringIO()
        self.skip_tags = {'script', 'style', 'head', 'meta', 'link'}
        self.current_tag = None
        self.block_tags = {'p', 'div', 'br', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 
                          'li', 'tr', 'td', 'th', 'section', 'article'}
    
    def handle_starttag(self, tag, attrs):
        self.current_tag = tag.lower()
        if tag.lower() in self.block_tags:
            self.result.write('\n')
    
    def handle_endtag(self, tag):
        if tag.lower() in self.block_tags:
            self.result.write('\n')
        self.current_tag = None
    
    def handle_data(self, data):
        if self.current_tag not in self.skip_tags:
            text = data.strip()
            if text:
                self.result.write(text + ' ')
    
    def get_text(self):
        # Clean up multiple newlines and spaces
        text = self.result.getvalue()
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        return text.strip()


class DocumentService:
    def __init__(self, db: Session):
        self.db = db
    
    def get_all(self) -> list[Document]:
        return self.db.query(Document).order_by(Document.upload_date.desc()).all()
    
    def get_by_id(self, document_id: int) -> Document:
        document = self.db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise NotFoundError("Document", document_id)
        return document
    
    def get_content(self, document_id: int) -> str | None:
        document = self.get_by_id(document_id)
        return document.content
    
    async def upload(self, file: UploadFile) -> Document:
        ext = file.filename.split(".")[-1].lower()
        unique_filename = f"{uuid.uuid4().hex}.{ext}"
        file_path = settings.upload_dir / unique_filename
        
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        document = Document(
            filename=unique_filename,
            original_filename=file.filename,
            file_type=ext,
            file_path=str(file_path),
            status="processing"
        )
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        
        try:
            text_content, page_count = self._extract_text(str(file_path), ext)
            document.content = text_content
            document.page_count = page_count
            document.status = "ready"
        except Exception as e:
            logger.error(f"Text extraction failed for {file.filename}: {e}")
            document.status = "error"
            document.error_message = str(e)
        
        self.db.commit()
        self.db.refresh(document)
        return document
    
    def _extract_text(self, file_path: str, file_type: str) -> tuple[str, int]:
        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type == "docx":
            return self._extract_docx(file_path)
        elif file_type == "txt":
            return self._extract_txt(file_path)
        elif file_type in ("html", "htm"):
            return self._extract_html(file_path)
        raise FileUploadError(f"Unsupported file type: {file_type}")
    
    def _extract_txt(self, file_path: str) -> tuple[str, int]:
        """Extract text from plain text files (useful for testing)."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        # Estimate 1 page per ~3000 characters
        page_count = max(1, len(content) // 3000)
        return content, page_count
    
    def _extract_pdf(self, file_path: str) -> tuple[str, int]:
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {i + 1}]\n{text}")
        return "\n\n".join(text_parts), page_count
    
    def _extract_docx(self, file_path: str) -> tuple[str, int]:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        page_count = max(1, len(paragraphs) // 30)
        return "\n\n".join(paragraphs), page_count
    
    def _extract_html(self, file_path: str) -> tuple[str, int]:
        """Extract text from HTML files (legal filings often come in HTML)."""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            # Fallback: read as binary and decode with errors ignored
            with open(file_path, 'rb') as f:
                content = f.read().decode('utf-8', errors='ignore')
        
        # Parse HTML and extract text
        parser = HTMLTextExtractor()
        parser.feed(content)
        text = parser.get_text()
        
        # Estimate pages (3000 chars per page)
        page_count = max(1, len(text) // 3000)
        return text, page_count
    
    def delete(self, document_id: int) -> None:
        document = self.get_by_id(document_id)
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
        self.db.delete(document)
        self.db.commit()
